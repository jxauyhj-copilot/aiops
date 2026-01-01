import os
import shutil
from config.settings import settings
from langchain_community.vectorstores import FAISS
from langchain_huggingface import HuggingFaceEmbeddings
from langchain.text_splitter import CharacterTextSplitter
from langchain.docstore.document import Document

DOCS_DIR = os.path.join(os.path.dirname(__file__), "documents")

def get_uploaded_documents():
    """
    Returns a list of filenames currently in the documents directory.
    """
    if not os.path.exists(DOCS_DIR):
        return []
    return [f for f in os.listdir(DOCS_DIR) if not f.startswith('.')]

def remove_document(filename: str):
    """
    Removes a document from the directory and rebuilds the index.
    """
    file_path = os.path.join(DOCS_DIR, filename)
    if os.path.exists(file_path):
        os.remove(file_path)
        ingest_documents() # Rebuild index
        return True
    return False

def ingest_documents():
    """
    Reads ALL documents from the documents/ directory, 
    creates embeddings, and saves the FAISS index.
    """
    print("--- Starting Knowledge Base Ingestion ---")
    
    documents = []

    # Load from Disk (User Uploads & System Docs)
    if not os.path.exists(DOCS_DIR):
        os.makedirs(DOCS_DIR)
    
    for filename in os.listdir(DOCS_DIR):
        if filename.startswith('.'): continue
        
        file_path = os.path.join(DOCS_DIR, filename)
        try:
            text = ""
            if filename.endswith(".txt"):
                with open(file_path, "r", encoding="utf-8") as f:
                    text = f.read()
            elif filename.endswith(".pdf"):
                from pypdf import PdfReader
                reader = PdfReader(file_path)
                for page in reader.pages:
                    text += page.extract_text() + "\n"
            elif filename.endswith(".docx"):
                import docx
                doc = docx.Document(file_path)
                for para in doc.paragraphs:
                    text += para.text + "\n"
            
            if text.strip():
                documents.append(Document(page_content=text, metadata={"source": filename}))
                print(f"Loaded: {filename}")
        except Exception as e:
            print(f"Error loading {filename}: {e}")

    # Splitter
    text_splitter = CharacterTextSplitter(chunk_size=500, chunk_overlap=50)
    texts = text_splitter.split_documents(documents)
    
    if not texts:
        print("No documents to ingest.")
        return

    # Embeddings
    print("Loading Local Embedding Model (all-MiniLM-L6-v2)...")
    embeddings = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")
    
    # Vector Store
    print("Creating Vector Store...")
    db = FAISS.from_documents(texts, embeddings)
    
    # Save
    save_path = settings.VECTOR_DB_PATH
    if not os.path.exists(os.path.dirname(save_path)):
        os.makedirs(os.path.dirname(save_path))
        
    db.save_local(save_path)
    print(f"--- Ingestion Complete. Index saved to {save_path} ---")

def add_document(file_content: bytes, file_type: str, filename: str = "uploaded_file"):
    """
    Saves the file to disk and triggers a full re-ingestion.
    """
    if not os.path.exists(DOCS_DIR):
        os.makedirs(DOCS_DIR)
        
    save_path = os.path.join(DOCS_DIR, filename)
    
    with open(save_path, "wb") as f:
        f.write(file_content)
    
    # Trigger re-ingestion to update the index
    try:
        ingest_documents()
        return "Document added and index rebuilt successfully."
    except Exception as e:
        return f"Error during ingestion: {e}"

if __name__ == "__main__":
    ingest_documents()
